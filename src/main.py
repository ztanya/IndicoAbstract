#!/usr/bin/env python3.7

from __future__ import unicode_literals

import string

import lxml.etree
import xml.etree.cElementTree as ET
from docxtpl import DocxTemplate
import getterConference
import getterAbstract
import arabic_roman
import AbstractClass
from docx import Document

from collections import defaultdict

# from docx.shared import Pt, RGBColor
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from pprint import pprint as pp

if __name__ == "__main__":

    conferenceInfoXml_filename = "../xml_sources/conference_info_boa.xml"

    tree_conference = ET.parse(conferenceInfoXml_filename)
    root_conference = tree_conference.getroot()
    getter_conference = getterConference.getterConference

    name_en = getter_conference.getConference(root_conference).name_en
    name_ru = getter_conference.getConference(root_conference).name_ru
    add_info_en = getter_conference.getConference(root_conference).addInfo_en
    add_info_ru = getter_conference.getConference(root_conference).addInfo_ru
    number = getter_conference.getConference(root_conference).number
    roman_number = arabic_roman.arabic_roman(number)
    year = getter_conference.getConference(root_conference).year
    date_en = getter_conference.getConference(root_conference).date_en
    date_ru = getter_conference.getConference(root_conference).date_ru
    desc_en = getter_conference.getConference(root_conference).desc_en
    desc_ru = getter_conference.getConference(root_conference).desc_ru

    doc = DocxTemplate("../tpl.docx")
    context = {'title_name_en': name_en.upper(),
               'title_name_ru': name_ru.upper(),
               'name_en': name_en,
               'name_ru': name_ru,
               'add_info_en': add_info_en,
               'number': number,
               'roman_number': roman_number,
               'year': year,
               'date_en': date_en,
               'date_ru': date_ru,
               'desc_en': desc_en,
               'desc_ru': desc_ru}

    doc.render(context)
    doc.save("../doc_results/generated_doc_boa.docx")

    document = Document("../doc_results/generated_doc_boa.docx")
    styles = document.styles
    count = 0

    abstractsXml_filename = "../xml_sources/Abstracts_boa.xml"

    tree_abstracts = ET.parse(abstractsXml_filename)
    root_abstracts = tree_abstracts.getroot()
    doc_abstracts = lxml.etree.parse(abstractsXml_filename)
    count_abstracts = doc_abstracts.xpath('count(//abstract)')
    getter_abstract = getterAbstract.getterAbstract

    abstracts = []

    for i in range(1, int(count_abstracts)):
        abstract = AbstractClass.Abstract()
        getted_ab = getter_abstract.getAbstract(root_abstracts, i)
        abstract.title = getted_ab.title
        abstract.content = getted_ab.content
        abstract.authors = getted_ab.authors
        abstract.track = getted_ab.track
        abstracts.append(abstract)


    # создать список всех Tracks, удалить повторяющиеся
    tracks = []
    for a in abstracts:
        if a.track not in tracks:
            tracks.append(a.track)

    # разбить все abstracts на группы по Tracks
    d = defaultdict(list)
    for a in abstracts:
        for t in tracks:
            if a.track == t:
                d[t].append(a)

    #print(d)

    def byName_key(person):
        return person.firstName.capitalize()

    for section in d.keys():
        if section != '':
            index = section.find('.')
            p = document.add_paragraph(style=styles["GRID_Title"])
            p.add_run(section[index+2:])
            document.add_page_break()
        for abstract in d.get(section):
            all_affiliations = []
            all_authors = abstract.authors
            all_authors = sorted(all_authors, key=byName_key)
            for person in all_authors:
                all_affiliations.append(person.affiliation)


            #создаем упорядоченное множество неповторяющихся affiliations
            all_affiliations_nonrepeat = []
            for item in all_affiliations:
                if item not in all_affiliations_nonrepeat:
                    all_affiliations_nonrepeat.append(item)

            #пишем заголовок абстракта
            p = document.add_paragraph(style=styles["GRID_Title"])
            p.add_run(abstract.title.upper().strip())

            # пишем авторов абстракта
            p = document.add_paragraph(style=styles["GRID_author"])
            aff_index = -1  # affiliation автора
            email_index = 0  # индекс буквы для email
            # генерируем имена  с индексами
            for i in range(0, len(all_authors)):
                if i > 0:
                    p.add_run(", ")
                p.add_run(all_authors[i].firstName.capitalize() + "\xa0" + all_authors[i].familyName.capitalize())  # имя + фамилия
                if len(all_affiliations_nonrepeat) > 1:
                    for index, elem in enumerate(all_affiliations_nonrepeat):
                        if all_authors[i].affiliation == elem:
                            aff_index = index+1
                            break
                    p.add_run(str(aff_index) + ",").font.superscript = True
                if all_authors[i].isPrimaryAuthor == True:
                    if all_authors[i].email != "":
                        p.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                        email_index += 1

            # пишем affiliations
            aff_index = -1
            affiliation = ""
            typed_affiliations = []  # affiliations, которые мы на странице уже написали (чтобы не повторяться)
            for person in all_authors:
                if (person.affiliation != "") and not (person.affiliation in typed_affiliations):
                    for index, elem in enumerate(all_affiliations_nonrepeat):
                        if person.affiliation == elem: #ищем совпадающий элемент во множестве
                            aff_index = index + 1 #запоминаем индекс элемента
                            affiliation = elem #запоминаем имя элемента
                            break
                    p = document.add_paragraph(style=styles["GRID_affiliation"])
                    if len(all_affiliations_nonrepeat) > 1:
                        p.add_run(str(aff_index)).font.superscript = True #пишем индекс
                    p.add_run(affiliation) #пишем название affiliation
                    typed_affiliations.append(affiliation) #добавляем affiliation в список уже написанных на странице


            # пишем e-mail
            count_primaryAuthors = 0
            p = document.add_paragraph(style=styles["GRID_email"])
            p.add_run("E-mail: ")
            email_index = 0  # буква для email
            for i in range(0, len(all_authors)):
                if all_authors[i].isPrimaryAuthor == True:
                    if all_authors[i].email != "":
                        if count_primaryAuthors > 0:
                            if i > 0:
                                p.add_run(", ")
                        p.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                        email_index += 1
                        p.add_run(all_authors[i].email)
                        count_primaryAuthors += 1

            # else:
            #     for person in all_authors:
            #         # генерируем имя
            #         p.add_run(person.firstName.capitalize()+ "\xa0" + person.familyName.capitalize())  # имя + фамилия
            #
            #         p = document.add_paragraph(style=styles["GRID_affiliation"])
            #         p.add_run(all_affiliations_nonrepeat[0])
            #
            #         p = document.add_paragraph(style=styles["GRID_email"])
            #         p.add_run("E-mail: " + person.email)

            p = document.add_paragraph(style=styles["GRID_Abstract"])
            p.add_run(abstract.content[1:-1].capitalize())
            document.add_paragraph("")

            document.add_page_break()

document.save("../doc_results/generated_final_boa.docx")
print("Документ успешно сгенерирован (../doc_results/generated_final_boa.docx)")
