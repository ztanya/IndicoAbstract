# -*- coding: utf-8 -*-
from __future__ import unicode_literals

import string
from collections import defaultdict
from os import path

from docxtpl import DocxTemplate
from docx import Document
from src.arabic_roman import arabic_roman


def get_list_of_tracks(list_abstracts):
    """Creates a list of all Tracks, removes duplicates."""
    tracks_list = []
    # First collect all possible tracks with repetition
    for abstract in list_abstracts:
        tracks_list.append(abstract.track)
    # Remove repetitions
    tracks_list = list(set(tracks_list))
    return tracks_list


def get_abstracts_by_tracks(list_abstracts, list_tracks):
    """Breaks all list_abstracts into groups by Tracks in list_tracks."""
    dict_abstracts_by_groups = defaultdict(list)
    for abstract in list_abstracts:
        for track in list_tracks:
            if abstract.track == track:
                dict_abstracts_by_groups[track].append(abstract)
    return dict_abstracts_by_groups


def by_name_key(person):
    """Returns person's first name with a capital."""
    return person.first_name.capitalize()


def generate_book(conference_obj, list_abstracts, doctpl_filename, finaldocument_filename):
    """The main function for generation of the book.
            Creates the final generated document.
            Includes the following steps:
            1. Parsing conference information XML file.
            2. Creating and rendering DOCX file from template for generation of
            two first book's pages.
            3. Saving the DOCX file as a temporary.
            4. Specifying about using styles in the generated document is the same as in the template.
            5. Parsing abstracts information XML file.
            6. Adding Abstracts class objects to list.
            7. Creating a list of all Tracks, removing duplicates.
            8. Grouping abstracts list by Tracks.
            9. Adding presentation themes (Tracks) to the final document.
            10. Creating an ordered set of all affiliations.
            11. Adding titles of the abstracts.
            12. Adding the abstracts author names with indexes of e-mails and
            indexes of affiliations to the final document.
            13. Adding author e-mails and affiliations to the final document.
            14. Saving the final generated document.
            15. Printing message about success generation and path to the generated document.
            conferenceinfo_xmlfilename - input path of XML file about information of the conference.
            abstracts_xmlfilename - input path of XML file about information of the abstracts.
            doctpl_filename - input path of DOCX template.
            finaldocument_filename - input path of DOCX generated document.
            """
    context = {'title_name_en': conference_obj.name_en.upper(),
               'title_name_ru': conference_obj.name_ru.upper(),
               'name_en': conference_obj.name_en,
               'name_ru': conference_obj.name_ru,
               'add_info_en': conference_obj.add_info_en,
               'add_info_ru': conference_obj.add_info_ru,
               'conf_number': conference_obj.number,
               'roman_number': arabic_roman(conference_obj.number),
               'year': conference_obj.year,
               'date_en': conference_obj.date_en,
               'date_ru': conference_obj.date_ru,
               'desc_en': conference_obj.desc_en,
               'desc_ru': conference_obj.desc_ru}
    doc = DocxTemplate(doctpl_filename)
    doc.render(context)
    tmp_doc = 'doc_results/generated_doc_tmp.docx'
    doc.save(tmp_doc)

    document = Document(tmp_doc)
    styles = document.styles

    tracks = get_list_of_tracks(list_abstracts)
    abstracts_by_tracks = get_abstracts_by_tracks(list_abstracts, tracks)

    for section in abstracts_by_tracks.keys():
        # Writes sections in the final document
        if section != '' and section != ' ' and section != None:
            # index = section.find('.')
            new_paragraph = document.add_paragraph(style=styles["GRID_Title"])
            # new_paragraph.add_run(section[index+2:])
            new_paragraph.add_run(section)
            document.add_page_break()
        for abstract in abstracts_by_tracks[section]:
            all_affiliations = []
            all_authors = abstract.authors
            all_authors = sorted(all_authors, key=by_name_key)
            for person in all_authors:
                all_affiliations.append(person.affiliation)

            # Creates an ordered set of non-repeating affiliations.
            all_affiliations_nonrepeat = []
            for item in all_affiliations:
                if item not in all_affiliations_nonrepeat:
                    all_affiliations_nonrepeat.append(item)

            # Writes abstract.Title in the final document
            new_paragraph = document.add_paragraph(style=styles["GRID_Title"])
            new_paragraph.add_run(abstract.title.upper().strip())

            # Writes abstract.Authors in the final document
            new_paragraph = document.add_paragraph(style=styles["GRID_author"])
            aff_index = -1  # affiliation автора
            email_index = 0  # индекс буквы для email
            # Generates names with indexes
            for i in range(0, len(all_authors)):
                if i > 0:
                    new_paragraph.add_run(", ")
                # FirstName + FamilyName:
                new_paragraph.add_run(all_authors[i].first_name.capitalize() + "\xa0" +
                                      all_authors[i].family_name.capitalize())
                if len(all_affiliations_nonrepeat) > 1:
                    for index, elem in enumerate(all_affiliations_nonrepeat):
                        if all_authors[i].affiliation == elem:
                            aff_index = index + 1
                            break
                    new_paragraph.add_run(str(aff_index)).font.superscript = True
                if all_authors[i].is_primary_author:
                    if all_authors[i].email != "":
                        if len(all_affiliations_nonrepeat) > 1:
                            new_paragraph.add_run("," +
                                                  string.ascii_lowercase[email_index]).font.superscript = True
                            email_index += 1
                        else:
                            new_paragraph.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                            email_index += 1

            # Writes abstract.Affiliations
            aff_index = -1
            affiliation = ""
            # Affiliations that we have already written (not to be repeated)
            typed_affiliations = []
            for person in all_authors:
                if (person.affiliation != "") and person.affiliation not in typed_affiliations:
                    for index, elem in enumerate(all_affiliations_nonrepeat):
                        if person.affiliation == elem:  # Looking for a matching element in the set
                            aff_index = index + 1  # Remember the index of the element
                            affiliation = elem  # Remember the name of the element
                            break
                    new_paragraph = document.add_paragraph(style=styles["GRID_affiliation"])
                    if len(all_affiliations_nonrepeat) > 1:
                        new_paragraph.add_run(
                            str(aff_index)).font.superscript = True  # write the index
                    new_paragraph.add_run(affiliation)  # write the affiliation
                    # add affiliation to the list of already written on the page:
                    typed_affiliations.append(affiliation)

            # write the e-mail
            count_primary_authors = 0
            new_paragraph = document.add_paragraph(style=styles["GRID_email"])
            new_paragraph.add_run("E-mail: ")
            email_index = 0  # letter for the email
            for i in range(0, len(all_authors)):
                if all_authors[i].is_primary_author:
                    if all_authors[i].email != "":
                        if count_primary_authors > 0:
                            if i > 0:
                                new_paragraph.add_run(", ")
                        new_paragraph.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                        email_index += 1
                        new_paragraph.add_run(all_authors[i].email)
                        count_primary_authors += 1

            #  write the content
            for x in abstract.content.split(chr(10)):
                if x != ' ' and x != '':
                    new_paragraph = document.add_paragraph(style=styles["GRID_Abstract"])
                    new_paragraph.add_run(x.replace(chr(9), ''))

            document.add_page_break()

    document.save(finaldocument_filename)
    print("The document has been successfully generated to the path: " + path.normpath(finaldocument_filename) + ".")
