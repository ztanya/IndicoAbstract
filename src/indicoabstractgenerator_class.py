"""Module for generating Book of abstracts."""

from __future__ import unicode_literals
import string
import xml.etree.cElementTree as ET
from collections import defaultdict
import lxml.etree as LXML_ET
from docxtpl import DocxTemplate
from docx import Document
import readline
import getter_conference_class
import getter_abstract_class
import arabic_roman
import abstract_class


class IndicoAbstractGenerator:
    """The main class for generation of the book."""

    def generate_doc(conferenceinfo_xmlfilename, abstracts_xmlfilename,
                     doctpl_filename, finaldocument_filename):
        """The main function for generation of the book.
        Creates the final generated document.
        Includes the following steps:
        1. Parsing conference information XML file.
        2. Creating and rendering DOCX file from template for generation of
        two first book's pages.
        3. Saving the DOCX file as a temporary.
        4. Specifying about using styles in the generated document is the same as in the template.
        5. Parsing abstracts information XML file.
        6. Adding Abstracts class objects to list.
        7. Creating a list of all Tracks, removing duplicates.
        8. Grouping abstracts list by Tracks.
        9. Adding presentation themes (Tracks) to the final document.
        10. Creating an ordered set of all affiliations.
        11. Adding titles of the abstracts.
        12. Adding the abstracts author names with indexes of e-mails and
        indexes of affiliations to the final document.
        13. Adding author e-mails and affiliations to the final document.
        14. Saving the final generated document.
        15. Printing message about success generation and path to the generated document.

        conferenceinfo_xmlfilename - input path of XML file about information of the conference.
        abstracts_xmlfilename - input path of XML file about information of the abstracts.
        doctpl_filename - input path of DOCX template.
        finaldocument_filename - input path of DOCX generated document.

        """

        tree_conference = ET.parse(conferenceinfo_xmlfilename)
        root_conference = tree_conference.getroot()
        getter_conference = getter_conference_class.GetterConference

        name_en = getter_conference.get_conference(root_conference).name_en
        name_ru = getter_conference.get_conference(root_conference).name_ru
        conf_number = getter_conference.get_conference(root_conference).number

        context = {'title_name_en': name_en.upper(),
                   'title_name_ru': name_ru.upper(),
                   'name_en': name_en,
                   'name_ru': name_ru,
                   'add_info_en': getter_conference.get_conference(root_conference).add_info_en,
                   'add_info_ru': getter_conference.get_conference(root_conference).add_info_ru,
                   'conf_number': getter_conference.get_conference(root_conference).number,
                   'roman_number': arabic_roman.arabic_roman(conf_number),
                   'year': getter_conference.get_conference(root_conference).year,
                   'date_en': getter_conference.get_conference(root_conference).date_en,
                   'date_ru': getter_conference.get_conference(root_conference).date_ru,
                   'desc_en': getter_conference.get_conference(root_conference).desc_en,
                   'desc_ru': getter_conference.get_conference(root_conference).desc_ru}

        doc = DocxTemplate(doctpl_filename)
        doc.render(context)
        doc.save("../doc_results/generated_doc_tmp.docx")

        document = Document("../doc_results/generated_doc_tmp.docx")
        styles = document.styles

        tree_abstracts = ET.parse(abstracts_xmlfilename)
        root_abstracts = tree_abstracts.getroot()
        doc_abstracts = LXML_ET.parse(abstracts_xmlfilename)
        count_abstracts = doc_abstracts.xpath('count(//abstract)')
        getter_abstracts = getter_abstract_class.GetterAbstract

        abstracts_list = []

        for i in range(1, int(count_abstracts)+1):
            getted_ab = getter_abstracts.get_abstract(root_abstracts, i)
            abstract = abstract_class.Abstract(getted_ab.title,
                                               getted_ab.content,
                                               getted_ab.authors,
                                               getted_ab.track)
            abstracts_list.append(abstract)

        # создать список всех Tracks, удалить повторяющиеся
        tracks_list = []
        for abstract in abstracts_list:
            if abstract.track not in tracks_list:
                tracks_list.append(abstract.track)

        # разбить все abstracts_list на группы по Tracks
        dict_abstracts_by_groups = defaultdict(list)
        for abstract in abstracts_list:
            for track in tracks_list:
                if abstract.track == track:
                    dict_abstracts_by_groups[track].append(abstract)

        def by_name_key(person):
            """Returns person's first name with a capital."""
            return person.first_name.capitalize()

        for section in dict_abstracts_by_groups.keys():
            if section != '':
                index = section.find('.')
                new_paragraph = document.add_paragraph(style=styles["GRID_Title"])
                new_paragraph.add_run(section[index + 2:])
                document.add_page_break()
            for abstract in dict_abstracts_by_groups.get(section):
                all_affiliations = []
                all_authors = abstract.authors
                all_authors = sorted(all_authors, key=by_name_key)
                for person in all_authors:
                    all_affiliations.append(person.affiliation)

                # создаем упорядоченное множество неповторяющихся affiliations
                all_affiliations_nonrepeat = []
                for item in all_affiliations:
                    if item not in all_affiliations_nonrepeat:
                        all_affiliations_nonrepeat.append(item)

                # пишем заголовок абстракта
                new_paragraph = document.add_paragraph(style=styles["GRID_Title"])
                new_paragraph.add_run(abstract.title.upper().strip())

                # пишем авторов абстракта
                new_paragraph = document.add_paragraph(style=styles["GRID_author"])
                aff_index = -1  # affiliation автора
                email_index = 0  # индекс буквы для email
                # генерируем имена  с индексами
                for i in range(0, len(all_authors)):
                    if i > 0:
                        new_paragraph.add_run(", ")
                    # имя + фамилия:
                    new_paragraph.add_run(all_authors[i].first_name.capitalize() + "\xa0" +
                                          all_authors[i].family_name.capitalize())
                    if len(all_affiliations_nonrepeat) > 1:
                        for index, elem in enumerate(all_affiliations_nonrepeat):
                            if all_authors[i].affiliation == elem:
                                aff_index = index + 1
                                break
                        new_paragraph.add_run(str(aff_index)).font.superscript = True
                    if all_authors[i].is_primary_author:
                        if all_authors[i].email != "":
                            if len(all_affiliations_nonrepeat) > 1:
                                new_paragraph.add_run(","+
                                    string.ascii_lowercase[email_index]).font.superscript = True
                                email_index += 1
                            else:
                                new_paragraph.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                                email_index += 1

                # пишем affiliations
                aff_index = -1
                affiliation = ""
                # affiliations, которые мы на странице уже написали (чтобы не повторяться)
                typed_affiliations = []
                for person in all_authors:
                    if (person.affiliation != "") and person.affiliation not in typed_affiliations:
                        for index, elem in enumerate(all_affiliations_nonrepeat):
                            if person.affiliation == elem:  # ищем совпадающий элемент во множестве
                                aff_index = index + 1  # запоминаем индекс элемента
                                affiliation = elem  # запоминаем имя элемента
                                break
                        new_paragraph = document.add_paragraph(style=styles["GRID_affiliation"])
                        if len(all_affiliations_nonrepeat) > 1:
                            new_paragraph.add_run(
                                str(aff_index)).font.superscript = True  # пишем индекс
                        new_paragraph.add_run(affiliation)  # пишем название affiliation
                        # добавляем affiliation в список уже написанных на странице:
                        typed_affiliations.append(affiliation)

                # пишем e-mail
                count_primary_authors = 0
                new_paragraph = document.add_paragraph(style=styles["GRID_email"])
                new_paragraph.add_run("E-mail: ")
                email_index = 0  # буква для email
                for i in range(0, len(all_authors)):
                    if all_authors[i].is_primary_author:
                        if all_authors[i].email != "":
                            if count_primary_authors > 0:
                                if i > 0:
                                    new_paragraph.add_run(", ")
                            new_paragraph.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                            email_index += 1
                            new_paragraph.add_run(all_authors[i].email)
                            count_primary_authors += 1

                #  пишем содержание
                for x in abstract.content.split(chr(10)):
                    if x != ' ' and x != '':
                        new_paragraph = document.add_paragraph(style=styles["GRID_Abstract"])
                        new_paragraph.add_run(x.replace(chr(9), ''))

                document.add_page_break()

        document.save(finaldocument_filename)
        print("Документ успешно сгенерирован (" + finaldocument_filename + ").")
