#!/usr/bin/env python3.7

from __future__ import unicode_literals

import string

import lxml.etree
import xml.etree.cElementTree as ET
from docxtpl import DocxTemplate
import getterConference
import getterAbstract
import arabic_roman
import AbstractClass
from docx import Document

from collections import defaultdict

# from docx.shared import Pt, RGBColor
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from pprint import pprint as pp

if __name__ == "__main__":

    tree_conference = ET.parse("conference_info.xml")
    root_conference = tree_conference.getroot()
    getter_conference = getterConference.getterConference

    name_en = getter_conference.getConference(root_conference).name_en
    name_ru = getter_conference.getConference(root_conference).name_ru
    add_info_en = getter_conference.getConference(root_conference).addInfo_en
    add_info_ru = getter_conference.getConference(root_conference).addInfo_ru
    number = getter_conference.getConference(root_conference).number
    roman_number = arabic_roman.arabic_roman(number)
    year = getter_conference.getConference(root_conference).year
    date_en = getter_conference.getConference(root_conference).date_en
    date_ru = getter_conference.getConference(root_conference).date_ru
    desc_en = getter_conference.getConference(root_conference).desc_en
    desc_ru = getter_conference.getConference(root_conference).desc_ru

    doc = DocxTemplate("tpl.docx")
    context = {'title_name_en': name_en.upper(),
               'title_name_ru': name_ru.upper(),
               'name_en': name_en,
               'name_ru': name_ru,
               'add_info_en': add_info_en,
               'number': number,
               'roman_number': roman_number,
               'year': year,
               'date_en': date_en,
               'date_ru': date_ru,
               'desc_en': desc_en,
               'desc_ru': desc_ru}

    doc.render(context)
    doc.save("generated_doc.docx")

    document = Document("generated_doc.docx")
    styles = document.styles
    count = 0

    tree_abstracts = ET.parse("Abstracts.xml")
    root_abstracts = tree_abstracts.getroot()
    doc_abstracts = lxml.etree.parse("Abstracts.xml")
    count_abstracts = doc_abstracts.xpath('count(//abstract)')
    getter_abstract = getterAbstract.getterAbstract_reformatted

    abstracts = []

    for i in range(1, int(count_abstracts)):
        abstract = AbstractClass.Abstract_reformatted()
        getted_ab = getter_abstract.getAbstract(root_abstracts, i)
        abstract.title = getted_ab.title
        abstract.content = getted_ab.content
        abstract.authors = getted_ab.authors
        abstract.track = getted_ab.track
        abstracts.append(abstract)


    # создать список всех Tracks, удалить повторяющиеся
    # разбить все abstracts на группы по Tracks
    # перед печатью в документ вставлять страницу с названием Track текущей группы abstracts

    tracks = []
    for a in abstracts:
        if a.track not in tracks:
            tracks.append(a.track)

    d = defaultdict(list)

    for a in abstracts:
        for t in tracks:
            if a.track == t:
                d[t].append(a.title)

    print(list(d.values()))

    """['\nГрид и облачная инфраструктура дата-центра Института Физики НАН Азербайджана\n',
     '\nClustering methods for energy consumption forecasting in smart grids\n',
     '\nClouds of JINR, University of Sofia and INRNE - current state of the project\n',
     '\nProblems of date and time data types in relational model of data\n',
     '\nMolecular dynamic simulation of water vapor interaction with various types of pores using hybrid computing structures\n',
     '\nElectronic, Dynamical and Thermodynamic Properties of DNA\n',
     '\nIMPROVING THE EFFICIENCY OF SMART GRIDS OF ENERGY CONSUMPTION WITH APPLICATION OF SYSTEMS OF ARTIFICIAL INTELLECT\n',
     '\nО методах и технологиях интеллектуального энергосбережения в коммерческих зданиях\n',
     '\nDeep machine learning and pattern/face recognition based on quantum neural networks and quantum genetic algorithm\n',
     '\nПрименение сети Хопфилда для автоматизированной подборки КПЭ\n',
     '\nA new approach to the development of provenance metadata management systems for large scientific experiments\n',
     '\nАнализ параллельной структуры популяционных алгоритмов оптимизации\n',
     '\nSensitivity Analysis in a problem of ReaxFF molecular-dynamic force field optimization\n',
     '\nCloud-based Computing for LHAASO experiment at IHEP\n',
     '\nOn porting of applications to new heterogeneous systems\n',
     '\nПрименение эволюционных и роевых алгоритмов оптимизации для решения модельной задачи предсказания структуры белка\n',
     '\nMULTIPARTICLE PRODUCTION, NEGATIVE BINOMIAL DISTRIBUTION, SUPERSYMMETRIC DYNAMICS, STATISTICAL SUMS AND ZETA-FUNCTIONS\n',
     '\nComparison of different convolution neural network architectures for the solution of the problem of emotion recognition by facial expression\n',
     '\nParallel calculations of ground states of 6,7,9,11Li nuclei by Feynman’s continual integrals method\n',
     '\nCombining satellite imagery and machine learning to predict atmospheric heavy metal contamination\n',
     '\nArchitecture and basic principles of the multifunctional platform for plant disease detection\n',
     '\nХаотическая динамика мгновенного сердечного ритма и его фазовое пространство.\n',
     '\nВизуализация квантового фазового пространства мгновенного сердечного ритма\n',
     'Text segmentation on photorealistic images', '\nIntegrating LEAF to data management workflow in LHAASO\n',
     '\nКогнитивно интеллектуальная система диагностики, адаптации и обучения детей-аутистов. Модуль обработки данных\n',
     '\nКогнитивно-интеллектуальная система адаптации и обучения детей-аутистов\n',
     '\nTHE USE OF SQUARE INTERPOLATION TO ACCELERATE THE CONVERGENCE OF THE CONTINUOUS ANALOGUE OF THE NEWTON METHOD\n',
     '\nBigPanDA Experience on Titan for the ATLAS Experiment at the LHC\n',
     '\nРазработка перспективной системы сбора данных на основе TRB-3\n',
     '\nEnabling Biology, Chemistry and Other Sciences on Titan through BigPanDA\n',
     '\nMulticomponent cluster management system for the computing center at IHEP\n',
     '\nA way of anomaly detection in engineering equipment characteristics of Symmetra at IHEP IT center\n',
     '\nEvent-Driven Automation and chat-ops on IHEP computing cluster\n',
     '\nEfficiency measurement system for the computing cluster at IHEP\n',
     '\nPseudo-random number generator based on neural network\n', 'THE BIGPANDA MONITORING SYSTEM ARCHITECTURE',
     '\nModernization of web service for the data center simulation program\n',
     '\nCluster analysis of scientific payload to execute it efficiently in distributed computing environment\n',
     '\nРеализация вычислений с динамическими зависимостями задач в среде десктоп грид с использованием Everest и Templet Web\n',
     '\nModeling of task scheduling in desktop grid systems at the initial stage of development\n',
     '\nPossible application areas of machine learning techniques at MPD/NICA experiment and their implementation prospects in distributed computing environment\n',
     '\nUsing binary file format description languages for verifying raw data in astroparticle physics experiments\n',
     '\nUsage of the distributed computing system in the recovery of the spectral density of sea waves\n',
     '\nDEVELOPMENT OF SOFTWARE FOR FACE RETRIEVAL SYSTEMS M ODELING\n',
     'Application of unified monitoring system in LHAASO',
     '\nSelection of rational composition of IT-services of information system with the purpose of increase of efficiency of transport logistics companies functioning\n',
     '\nOpenFOAM wave modelling optimization with heterogeneous systems application porting.\n',
     '\nИсследование скорости сбора информации в беспроводных сенсорных сетях.\n',
     '\nA distributed data warehouse system for astroparticle physics\n',
     '\nParticle identification in ground-based gamma-ray astronomy using convolutional neural networks\n',
     'Data Preprocessing for Credit Scoring',
     '\nDesign and implementation of a service for performing HPC computations in cloud environment.\n',
     '\nExperiments with JupyterHub at the Saint Petersburg State University\n',
     '\nParticipation of Russian institutions in the processing and storage of ALICE data.\n',
     '\nAPPROACHES TO THE AUTOMATED DEPLOYMENT OF THE CLOUD INFRASTRUCTURE OF GEOGRAPHICALLY DISTRIBUTED DATA CENTERS\n',
     '\nDirect Simulation of the Charge Transfer along Oligonucleotides at T=300K\n',
     'The distributed grid site of Institute of Physics',
     '\nAgent Technology Situational Express Analysis in Assessment of Technological Development Level of the BRICS Countries\n',
     '\nProperties of The Parallel Discrete Event Simulation Algorithms on Small-World Communication Networks\n',
     '\nСверточная нейронная сеть в системе стереозрения мобильного робота\n',
     '\nThe interoperability problem during the implementation of the FOURTH PARADIGM\n',
     '\nBOINC-based comparison of the geoacoustic inversion algorithms efficiency\n',
     '\nОРГАНИЗАЦИЯ ДОСТУПА К ЭКСПЕРИМЕНТАЛЬНЫМ ДАННЫМ УСТАНОВКИ ИТЭР В РЕЖИМЕ УДАЛЕННОЙ ПУЛЬТОВОЙ\n',
     '\nSemantic information management: the approach to semantic assets development lifecycle\n',
     '\nANALYSIS OF THE FEATURES OF THE OPTIMAL LOGICAL STRUCTURE OF DISTRIBUTED DATABASES\n',
     'Combined Explicit-Implicit Taylor Series Methods', 'INP BSU grid site',
     '\nJINR Multifunctional Information and Computing Complex: Status and Perspectives\n',
     '\nDevelopment of JINR Tier-1 service monitoring system\n', 'File Transfer Service at Exabyte scale',
     '\nThe problem of symbolic-numeric computation of the eigenvalues and eigenfunctions of the leaky modes in a regular homogeneous open waveguide\n',
     'Cloud Meta-Scheduler for Dynamic VM Reallocation',
     '\nNumerical solution of diffraction problem on the joint of two open three-layered waveguides\n',
     '\nIntegration of the BOINC system and additional software packages\n',
     '\nМногомерный анализ данных о продажах на основе технологии OLAP\n',
     '\nSupporting Efficient Execution of Many-Task Applications with Everest\n', 'New features of the JINR cloud',
     '\nThe BigPanDA self-monitoring alarm system for ATLAS\n',
     '\nCurrent status of software development for the MPD and BM@N experiments\n',
     '\nTier-1 centre at NRC «Kurchatov institute» between LHC Run2 and Run3\n',
     'Experience with ITEP-FRRC HPC facility', 'Data Knowledge Base for the ATLAS collaboration',
     '\nNew methods of minimizing the errors in the software\n',
     '\nИмитационная модель БРЛК с синтезированной апертурой антенны в сети распределенных вычислений MarGrid\n',
     '\nMechanisms for ensuring the integrity of information in distributed computing systems in the long-term period of time\n',
     '\nInformation-software environment for the GOVORUN supercomuter\n', 'HybriLIT monitoring system',
     '\nScalability of the Parallel Strongin Algorithm in the Problem of Optimizing a Molecular-Dynamic Force Field\n',
     '\nMerging multidimensional histograms via hypercube algorithm\n',
     '\nImproving the load of supercomputers based on job migration using container virtualization\n',
     '\nБиблиотеки и пакеты прикладных программ, доступные пользователям ЭВМ ОИЯИ\n',
     '\nИсследование особенностей Интернет-трафика в магистральном канале\n',
     '\nThe activity of Russian Chapter of International Desktop Grid Federation\n',
     '\nKubernetes testbed cluster for the Lightweight Sites project\n',
     '\nData gathering and wrangling for the monitoring of the Russian labour market\n',
     '\nTHE GAME CHARACTER OF COLLABORATION IN VOLUNTEER COMPUTING COMMUNITY\n',
     '\nDevelopment of the Geometry Database for the CBM Experiment and first adjustment for use in NICA project\n',
     'Performance measurements for the WLCG Cost Model',
     '\nThe concept of proactive protection in a distributed computing system\n']"""


    def byName_key(person):
        return person.firstName.capitalize()

    for abstract in abstracts:
        all_affiliations = []
        all_authors = abstract.authors
        all_authors = sorted(all_authors, key=byName_key)
        for person in all_authors:
            all_affiliations.append(person.affiliation)


        #создаем упорядоченное множество неповторяющихся affiliations
        all_affiliations_nonrepeat = []
        for item in all_affiliations:
            if item not in all_affiliations_nonrepeat:
                all_affiliations_nonrepeat.append(item)

        #пишем заголовок абстракта
        p = document.add_paragraph(style=styles["GRID_Title"])
        p.add_run(abstract.title.upper().strip())

        # пишем авторов абстракта
        p = document.add_paragraph(style=styles["GRID_author"])
        aff_index = -1  # affiliation автора
        email_index = 0  # индекс буквы для email
        # генерируем имена  с индексами
        for i in range(0, len(all_authors)):
            if i > 0:
                p.add_run(", ")
            p.add_run(all_authors[i].firstName.capitalize() + "\xa0" + all_authors[i].familyName.capitalize())  # имя + фамилия
            if len(all_affiliations_nonrepeat) > 1:
                for index, elem in enumerate(all_affiliations_nonrepeat):
                    if all_authors[i].affiliation == elem:
                        aff_index = index+1
                        break
                p.add_run(str(aff_index) + ",").font.superscript = True
            if all_authors[i].email != "":
                p.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                email_index += 1

        # пишем affiliations
        aff_index = -1
        affiliation = ""
        typed_affiliations = []  # affiliations, которые мы на странице уже написали (чтобы не повторяться)
        for person in all_authors:
            if (person.affiliation != "") and not (person.affiliation in typed_affiliations):
                for index, elem in enumerate(all_affiliations_nonrepeat):
                    if person.affiliation == elem: #ищем совпадающий элемент во множестве
                        aff_index = index + 1 #запоминаем индекс элемента
                        affiliation = elem #запоминаем имя элемента
                        break
                p = document.add_paragraph(style=styles["GRID_affiliation"])
                if len(all_affiliations_nonrepeat) > 1:
                    p.add_run(str(aff_index)).font.superscript = True #пишем индекс
                p.add_run(affiliation) #пишем название affiliation
                typed_affiliations.append(affiliation) #добавляем affiliation в список уже написанных на странице


        # пишем e-mail
        p = document.add_paragraph(style=styles["GRID_email"])
        p.add_run("E-mail: ")
        email_index = 0  # буква для email
        for i in range(0, len(all_authors)):
            if all_authors[i].email != "":
                if i > 0:
                    p.add_run(", ")
                p.add_run(string.ascii_lowercase[email_index]).font.superscript = True
                email_index += 1
                p.add_run(all_authors[i].email)

    # else:
    #     for person in all_authors:
    #         # генерируем имя
    #         p.add_run(person.firstName.capitalize()+ "\xa0" + person.familyName.capitalize())  # имя + фамилия
    #
    #         p = document.add_paragraph(style=styles["GRID_affiliation"])
    #         p.add_run(all_affiliations_nonrepeat[0])
    #
    #         p = document.add_paragraph(style=styles["GRID_email"])
    #         p.add_run("E-mail: " + person.email)

        p = document.add_paragraph(style=styles["GRID_Abstract"])
        p.add_run(abstract.content[1:-1].capitalize())
        document.add_paragraph("")

        document.add_page_break()

document.save("generated_final.docx")
print("Документ успешно сгенерирован (generated_final.docx)")
